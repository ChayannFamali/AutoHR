import logging
import os
from typing import Optional, Tuple

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)

class DocumentParser:
    """Класс для парсинга PDF и DOCX документов"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, dict]:
        """
        Извлекает текст из PDF файла
        
        Returns:
            tuple: (extracted_text, metadata)
        """
        try:
            doc = fitz.open(file_path)
            text = ""
            metadata = {
                'pages': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'creator': doc.metadata.get('creator', ''),
            }
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
                text += "\n\n"  # Разделитель страниц
            
            doc.close()
            
            logger.info(f"Successfully extracted text from PDF: {file_path}")
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return "", {}
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, dict]:
        """
        Извлекает текст из DOCX файла
        
        Returns:
            tuple: (extracted_text, metadata)
        """
        try:
            doc = Document(file_path)
            text = ""
            metadata = {
                'paragraphs': 0,
                'tables': len(doc.tables),
            }
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    metadata['paragraphs'] += 1
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return "", {}
    
    @classmethod
    def parse_document(cls, file_path: str) -> Tuple[str, dict]:
        """
        Определяет тип файла и извлекает текст
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            tuple: (extracted_text, metadata)
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return "", {}
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return cls.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            return "", {}
    
    @staticmethod
    def validate_extracted_text(text: str) -> dict:
        """
        Валидирует извлеченный текст и возвращает метрики качества
        
        Returns:
            dict: Метрики качества текста
        """
        if not text or not text.strip():
            return {
                'is_valid': False,
                'error': 'No text extracted',
                'word_count': 0,
                'char_count': 0
            }
        
        words = text.split()
        
        return {
            'is_valid': True,
            'word_count': len(words),
            'char_count': len(text),
            'has_email': '@' in text,
            'has_phone': any(char.isdigit() for char in text),
            'avg_word_length': sum(len(word) for word in words) / len(words) if words else 0,
            'line_count': len(text.split('\n'))
        }
